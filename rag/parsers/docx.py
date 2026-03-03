"""DOCX parser using python-docx."""

from __future__ import annotations


def parse(path: str) -> str:
    """
    Extract text from a .docx Word document.

    Args:
        path: Path to the .docx file.

    Returns:
        Extracted plain text with paragraph breaks.
    """
    try:
        from docx import Document  # type: ignore[import-untyped]

        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except ImportError:
        return "[python-docx not installed. Run: pip install python-docx]"
    except Exception as exc:
        return f"[DOCX parse error: {exc}]"
